"""
Generates: KARIOS_MBSE_Process.docx
Detailed & referenced Model-Based Systems Engineering (MBSE) process for the
KARIOS CubeSat, grounded in the mission draft and the INCOSE SSWG CubeSat
Reference Model (CRM) / RAX MBSE work (Kaslow et al.), with the schematic
block diagram embedded.

Run:  python build_mbse_doc.py
Requires: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "KARIOS_MBSE_Process.docx"
DIAGRAM_PNG = "_mbse_preview-1.png"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
INK2   = RGBColor(0x3F, 0x4B, 0x5E)
C_REQ  = RGBColor(0x2F, 0x74, 0xC0)
C_STR  = RGBColor(0x1F, 0x8F, 0x86)
C_BEH  = RGBColor(0x7C, 0x53, 0xA5)
C_PAR  = RGBColor(0xC9, 0x76, 0x1F)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def make_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ============================================================ TITLE
tp = doc.add_heading("KARIOS CubeSat — Model-Based Systems Engineering (MBSE) Process", level=0)
for r in tp.runs:
    r.font.color.rgb = ACCENT
para("An MBSE approach for the KAIST Intelligence & Reconnaissance Orbital Satellite (KARIOS), "
     "instantiated from the INCOSE Space Systems Working Group (SSWG) CubeSat Reference Model "
     "(CRM) and adapted from the RANDEV CubeSat programme.", italic=True)

# ============================================================ 1
h("1. Purpose and Rationale", 1)
para("Small satellites such as KARIOS are challenging to design: they have tightly constrained "
     "resources, strongly coupled subsystems, and must operate in a dynamic space environment. "
     "Model-Based Systems Engineering (MBSE) addresses this by building a single, authoritative "
     "system model that integrates requirements, architecture, behaviour, and analysis in one "
     "coherent source of truth.")
para("Per the INCOSE MBSE Initiative (Systems Engineering Vision 2020), the KARIOS system model "
     "is initiated at the start of the project and evolves throughout development and into "
     "operations, providing a consistent basis for requirements, design, analysis, verification, "
     "and validation. The approach follows the INCOSE SSWG CubeSat Reference Model — the same "
     "lineage as the FireSat and Radio Aurora Explorer (RAX) reference models (Kaslow et al.) — "
     "and is tailored to KARIOS from the RANDEV CubeSat MBSE description.")

# ============================================================ 2
h("2. CubeSat Domain and Enterprise (Context)", 1)
para("Following the CRM, the KARIOS model is framed within the CubeSat Domain, which comprises "
     "the Mission Enterprise, its Stakeholders, the External Environment, and External Constraints.")
make_table(
    ["Domain element", "KARIOS instantiation"],
    [
        ["Mission Enterprise", "Development, deployment, and operation of the KARIOS SIGINT + "
         "cybersecurity CubeSat mission by KAIST ASCL."],
        ["Stakeholders & Concerns", "KAIST ASCL; IITP (funding, AI-satellite-vulnerability "
         "project); spectrum & regulatory agencies; launch/deployer provider; Cal Poly CubeSat "
         "Design Specification; INCOSE / OMG (SysML)."],
        ["External Environment", "Space environment (radiation, thermal, vacuum, debris) and "
         "Earth environment (ground stations, atmosphere)."],
        ["External Constraints", "Licenses & regulations — radio-frequency spectrum authorisation "
         "and orbital-debris / disposal guidelines."],
    ],
)
para("Segments are modelled at a common level: the Space Segment (KARIOS spacecraft), the Ground "
     "Segment (UHF and LoRaWAN ground-station services), and Transport / Launch / Deploy services "
     "procured from an external provider.", italic=True)

# ============================================================ 3
h("3. Stakeholders, Concerns, Viewpoints and Views", 1)
para("Per ISO/IEC/IEEE 42010, the model organises stakeholder concerns into viewpoints, and "
     "renders views that frame specific concerns. For KARIOS this means:")
bullet("Analyse mission requirements to identify enterprise-level use cases.")
bullet("Define the relationship between requirements and those use cases.")
bullet("Develop and capture use-case descriptions and scenarios in the model.")
bullet("Establish architecture viewpoints and render architecture views (structure, behaviour) "
       "that address regulatory, operational, and design concerns.")
para("Regulatory viewpoints (spectrum authorisation, orbital-debris compliance) are captured "
     "explicitly, since licenses, timelines, and procedures must be well understood and traceable "
     "within the model.")

# ============================================================ 4
h("4. The KARIOS SysML System Model — Four Pillars", 1)
para("The system model is authored in SysML (MagicDraw / Cameo Systems Modeler) and structured "
     "around the four SysML pillars. Each pillar is instantiated with KARIOS content:")

p = doc.add_paragraph(); r = p.add_run("Requirements. "); r.bold = True; r.font.color.rgb = C_REQ
p.add_run("Mission requirements are decomposed to system and subsystem requirements. "
          "Satisfy relationships link requirements to the value properties they constrain, and "
          "Verify relationships link them to the verification actions that close them — forming "
          "the Requirements Verification Matrix (RVM).")

p = doc.add_paragraph(); r = p.add_run("Structure. "); r.bold = True; r.font.color.rgb = C_STR
p.add_run("Block Definition and Internal Block Diagrams (BDD / IBD) capture the architecture: "
          "Space, Ground, and Launch segments; the bus subsystems — OBC, EPS, ADCS, "
          "Communications (UHF + LoRaWAN), Structure, and Thermal Control; and the payloads — the "
          "SIGINT primary payload and the cybersecurity secondary payload — together with their "
          "interfaces.")

p = doc.add_paragraph(); r = p.add_run("Behaviour. "); r.bold = True; r.font.color.rgb = C_BEH
p.add_run("Use cases, activity diagrams, and state machines model operations. A top-level "
          "'Collect & Distribute Mission Data' activity decomposes into mission tasking, data "
          "collection, and downlink. A CubeSat state machine captures Orbit, Solar (sunlight / "
          "eclipse), Payload (collecting / idle), and Download (transmitting / not) states, which "
          "may overlap operationally.")

p = doc.add_paragraph(); r = p.add_run("Parametrics. "); r.bold = True; r.font.color.rgb = C_PAR
p.add_run("Parametric diagrams express the engineering budgets as constraint blocks — power "
          "(generation vs. load and battery state), data (buffer fill vs. downlink volume), the "
          "communications link budget, and coverage / contact time — evaluated at each simulation "
          "time step.")

# ============================================================ 5
h("5. Integrated Analyses — an Executable Model", 1)
para("The distinctive strength of the CRM/RAX approach is that the SysML model is executable: "
     "discipline analyses are linked to it and time-stepped so that behaviour and parametrics run "
     "as a simulation. KARIOS adopts the same integration pattern:")
make_table(
    ["Tool", "Role in the KARIOS model"],
    [
        ["STK", "Spacecraft orbit, coverage, ground-station contacts, and eclipse timing "
         "(the pre-launch orbital dynamics & coverage analysis)."],
        ["MATLAB", "Analytical models — power generation/consumption, thermal, and link budgets."],
        ["ModelCenter + MBSE Analyzer", "Bridge that imports integrated analyses into SysML as "
         "constraint blocks and executes parametric diagrams."],
        ["Cameo Simulation Toolkit", "Time-steps the behavioural models (activities & state "
         "machine), calling the engineering analyses at each iteration."],
    ],
)
para("Results flow back into the parametric diagrams, whose outputs satisfy — or violate — the "
     "requirements they are linked to. This closes the loop between design and analysis inside a "
     "single model, and lets operational trade studies (e.g. energy and data-collection capacity) "
     "be run directly on the model.")

# ============================================================ 6
h("6. Verification and Validation within the Model", 1)
para("The CRM is described as 'a model of a model': a reference model that a mission team "
     "instantiates. For KARIOS the instantiated model supports both:")
bullet("Validation — objective evidence that the right system is being built, i.e. that it "
       "satisfies stakeholder needs.")
bullet("Verification — objective evidence that the system is built right, i.e. that each element "
       "performs its intended function and meets its allocated requirements. Verification methods "
       "are Inspection, Analysis, Demonstration, and Test.")
para("Behavioural and parametric simulation results provide analysis evidence directly from the "
     "model; the results are recorded against the RVM and feed the KARIOS Verification Control "
     "Document and the broader Verification Process. Regulatory compliance (spectrum, debris) is "
     "verified through dedicated viewpoints into the model.")

# ============================================================ 7  DIAGRAM
h("7. Schematic Block Diagram of the MBSE Process", 1)
para("The diagram below summarises the KARIOS MBSE process: the CubeSat Domain context frames an "
     "iterative SysML method (steps 1–6) that builds the four-pillar system model; integrated "
     "analyses (STK, MATLAB) are linked to the model as an executable co-simulation; and the "
     "results drive verification & validation, feeding the Verification Process. Feedback from "
     "V&V refines the model iteratively.")
if os.path.exists(DIAGRAM_PNG):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Figure 1 — KARIOS MBSE process (rendered from KARIOS_mbse_block_diagram.tex).")
    rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = INK2
    pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(DIAGRAM_PNG, width=Inches(6.4))
else:
    para("[Diagram PNG not found — compile KARIOS_mbse_block_diagram.tex and rasterize it.]",
         italic=True)

# ============================================================ 8
h("8. References", 1)
for ref in [
    "D. Kaslow et al., \"Integrated Model-Based Systems Engineering (MBSE) Applied to the "
    "Simulation of a CubeSat Mission\" (RAX), IEEE Aerospace Conference.",
    "D. Kaslow et al., \"Developing a CubeSat Model-Based System Engineering (MBSE) Reference "
    "Model — Interim Status,\" INCOSE / IEEE.",
    "INCOSE Space Systems Working Group (SSWG) CubeSat Reference Model (CRM); OMG SysML "
    "specification.",
    "INCOSE Systems Engineering Vision 2020 — MBSE Initiative.",
    "ISO/IEC/IEEE 42010 — Architecture description (stakeholders, concerns, viewpoints, views).",
    "Cal Poly CubeSat Design Specification (CDS).",
    "KARIOS mission draft (KAIST Intelligence & Reconnaissance Orbital Satellite); RANDEV CubeSat "
    "MBSE description.",
]:
    doc.add_paragraph(ref, style="List Bullet")

doc.save(OUT)
print("Wrote", OUT, "(diagram embedded)" if os.path.exists(DIAGRAM_PNG) else "(no diagram PNG)")
