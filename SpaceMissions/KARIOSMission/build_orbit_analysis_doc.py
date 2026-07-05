"""
Generates: KARIOS_Orbit_Design_and_Analysis.docx
The "Orbital Dynamics and Coverage Analysis" section for the KARIOS CubeSat,
formatted for inclusion in the MBSE / verification document.

Run:  python build_orbit_analysis_doc.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "KARIOS_Orbit_Design_and_Analysis.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # dark blue

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def make_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ============================================================ TITLE
title = doc.add_heading("KARIOS CubeSat — Orbital Dynamics and Coverage Analysis", level=0)
for run in title.runs:
    run.font.color.rgb = ACCENT
para("Pre-Launch Orbit Design, Simulation, and Verification using STK (Systems Tool Kit)",
     italic=True)
para("Verification method: Analysis (A) — ref. ECSS-E-ST-10 life-cycle standard and the "
     "KARIOS MBSE verification framework. This section demonstrates, by analysis, that the "
     "KARIOS design orbit satisfies the mission coverage, lighting, contact, and disposal "
     "requirements across the realistic launch-injection envelope.")

# ============================================================ 1. REFERENCE ORBIT
h("1. Reference Orbit Definition", 1)
para("The KARIOS design orbit is derived from the CAS500-3 (차세대중형위성3호, NORAD ID 66655) "
     "rideshare heritage two-line element set. The nominal osculating/mean elements adopted "
     "for pre-launch simulation are:")

make_table(
    ["Element", "Value", "Notes"],
    [
        ["Orbit regime", "Sun-Synchronous (SSO), near-circular", "Frozen orbit design"],
        ["Mean altitude", "~550 km", "Derived from reference TLE"],
        ["Inclination", "97.7444 deg", "SSO-consistent at this altitude"],
        ["RAAN", "108.1576 deg", "Sets Local Time of Ascending Node (LTAN)"],
        ["Eccentricity", "0.0006955", "Near-circular"],
        ["Argument of Perigee", "79.2781 deg", "Frozen-orbit phasing"],
        ["Mean Anomaly", "280.9219 deg", "Epoch phasing"],
        ["Orbital Period", "96.75 min", ""],
        ["Mean Motion", "14.88 rev/day", ""],
        ["Heritage reference", "CAS500-3 TLE (2025-11-27 / 2026-06-29)", "Rideshare cluster reference"],
    ],
)

para("Critical validation: The SSO inclination for a 550 km circular orbit is approximately "
     "97.6 deg. The adopted 97.7444 deg originates from the reference TLE (slightly different "
     "mean altitude). Before any downstream analysis, the nodal precession rate must be "
     "confirmed in STK to equal +0.9856 deg/day (360 deg / 365.25 days). If it does not, the "
     "orbit is not truly sun-synchronous and all lighting and coverage conclusions are invalid.",
     bold=True)

# ============================================================ 2. SCENARIO
h("2. STK Scenario Setup", 1)
numbered("Create a new scenario. Set the analysis period to at least one full repeat-ground-"
         "track cycle plus margin (>= 31 days) to capture LTAN drift, eclipse-season variation, "
         "and coverage statistics. Use a separate multi-year scenario for lifetime/decay.")
numbered("Set the scenario epoch to the planned launch/commissioning epoch (not the raw TLE "
         "epoch) so lighting and RAAN geometry are physically correct at launch.")
numbered("Work in J2000 for elements; report ground tracks in Earth-fixed (ECEF). Time in UTCG.")
numbered("Central body: Earth, WGS84 ellipsoid.")

# ============================================================ 3. THREE DEFS
h("3. Build KARIOS — Three Parallel Orbit Definitions", 1)
para("Define the same orbit three ways and confirm convergence. Disagreement is the earliest "
     "and most reliable error detector.")

h("3.1  From the CAS500-3 TLE (heritage truth reference)", 2)
bullet("Insert Satellite -> Propagator: SGP4. Load NORAD 66655 TLE.")
bullet("Use the latest TLE (2026-06-29) for current state; launch-day TLE (2025-11-27) for "
       "as-deployed conditions. This is the ground truth for the cluster and phasing.")

h("3.2  From Keplerian elements (design orbit)", 2)
bullet("New satellite -> Propagator: J2Perturbation for a fast sanity check, then HPOP (Section 4).")
bullet("Enter the six elements from Table 1 using Mean (Brouwer-Lyddane) element type when "
       "transcribing a TLE-derived mean state. Do not paste TLE mean elements directly into an "
       "osculating-element field.")

h("3.3  From the SSO / Repeat-Ground-Track wizard (design intent)", 2)
bullet("Use Orbit Wizard -> Sun Synchronous (or Repeating Ground Trace) to let STK solve the "
       "true sun-synchronous inclination, then compare to 97.7444 deg. The delta reveals whether "
       "the TLE orbit is truly frozen-SSO or slightly off.")
para("Deliverable: a reconciliation table showing definitions A, B, C agree on altitude, "
     "period, inclination, and RAAN precession within tight tolerance.", italic=True)

# ============================================================ 4. PROPAGATOR
h("4. Propagator Fidelity (HPOP Configuration)", 1)
para("SGP4 is a mean-element model and is not sufficient for coverage/lighting truth. Use HPOP "
     "for the design satellite:")
bullet("Gravity: EGM2008 (or EGM96), degree/order >= 21x21 (70x70 for decay studies).")
bullet("Drag: ON (dominant at 550 km). Atmosphere: NRLMSISE-00 or Jacchia-Roberts. "
       "Cd ~ 2.2 (document assumption). Area/mass from actual CubeSat mass and cross-section; "
       "run min-drag (velocity-pointing narrow face) and max-drag (broadside) cases as a bracket.")
bullet("Solar flux F10.7 and geomagnetic Ap: use predicted values for the launch epoch "
       "(SpaceWeather file). For lifetime, run nominal / +2sigma / -2sigma flux.")
bullet("SRP: ON, Cr ~ 1.2-1.5, dual-cone shadow model.")
bullet("Third body: Sun + Moon. Solid tides optional at this altitude/duration.")
para("Validation rule: match HPOP output against the SGP4/TLE ephemeris over 7-14 days. "
     "Position agreement to a few km early on validates the force model and area/mass.", bold=True)

# ============================================================ 5. SSO CHECK
h("5. Sun-Synchronicity and LTAN Validation (Core Check)", 1)
numbered("Report RAAN vs. time; confirm nodal precession ~ +0.9856 deg/day.")
numbered("Compute LTAN using the lighting / Sun-beta-angle report. Plot beta angle over a full "
         "year; it drives eclipse duration and thermal/power.")
numbered("Identify eclipse-season extremes (max eclipse fraction) and any full-sun periods if "
         "LTAN approaches 06:00/18:00. Feeds power and thermal verification.")
numbered("Report Sun beta min/max and max eclipse minutes per orbit to the EPS/thermal analysts.")

# ============================================================ 6. COVERAGE
h("6. Coverage Analysis", 1)

h("6.1  Ground-Station / Mission Contacts", 2)
bullet("Create Facility/Place objects for the primary and backup ground stations "
       "(include global network stations if used).")
bullet("Add a Sensor with a realistic minimum elevation mask (5-10 deg UHF/S-band; 10 deg X-band).")
bullet("Compute Access (satellite <-> GS). Report per day: number of passes, durations, "
       "max elevation, and gap statistics (max/mean).")
bullet("Report total contact minutes/day against the required downlink data volume.")
bullet("Overlay a link budget (STK Comms, if licensed) to confirm passes are usable, "
       "not merely geometric.")

h("6.2  Payload Coverage (Imaging / Sensing)", 2)
bullet("Add a Sensor modeling the payload FOV (rectangular or conical; half-angle/swath from "
       "the optics/GSD spec).")
bullet("Point the sensor per the attitude mode (nadir or target/slewing) with body-fixed offsets.")
bullet("Use Coverage Definition + Figure of Merit over a grid (global, latitude band, or AOIs). "
       "FOMs: Revisit Time, Access Duration, % Coverage, Number of Accesses, N-asset coverage.")
bullet("Constrain by Sun elevation at target (optical imaging: daytime, sun-angle limits) "
       "and by eclipse (sensor off in shadow).")
bullet("Report revisit time and daily/area coverage, and days-to-full-area-mapping.")

h("6.3  Lighting / Quality Constraints", 2)
bullet("Add solar-elevation and sensor-to-Sun constraints so coverage counts only when "
       "imaging is physically useful.")

# ============================================================ 7. CLUSTER
h("7. Rideshare / Cluster and Collision Context", 1)
bullet("Load CAS500-3 and other cluster members (CelesTrak/Space-Track TLEs) alongside KARIOS.")
bullet("Compute relative range / along-track separation vs. time to understand post-deployment "
       "drift and initial acquisition geometry.")
bullet("Screen against the catalog (conjunction awareness) if STK CAT / Space-Track access is "
       "available; at minimum note the congested 550 km SSO shell.")
bullet("If KARIOS has no propulsion, document that phasing is deployment-fixed and passive "
       "differential drag is the only control authority.")

# ============================================================ 8. LIFETIME
h("8. Orbit Lifetime and Decay", 1)
bullet("Run the Lifetime tool (or long HPOP propagation) with Cd, area/mass (min/max attitude), "
       "and F10.7/Ap predictions with 1-2sigma solar-cycle uncertainty.")
bullet("Output: estimated orbital lifetime and altitude-vs-time.")
bullet("Compliance: confirm re-entry within the applicable rule (25-year, or the newer 5-year "
       "guideline / KARI-MSIT national licensing requirement). This is a licensing deliverable.")

# ============================================================ 9. LINKAGE
h("9. Attitude, Power, and Thermal Linkage (MBSE Feed)", 1)
para("STK produces the geometry the other analyses require. Export as named, regenerable reports "
     "to support model-based (SysML/MBSE) traceability:")
bullet("Beta angle, eclipse entry/exit/duration -> EPS battery DoD and thermal cycling.")
bullet("Sun vector in body frame (with attitude profile) -> solar-array power and thermal loads.")
bullet("Nadir/target pointing angles and slew rates -> ADCS sizing.")
bullet("GS access windows -> operations timeline and data budget.")

# ============================================================ 10. SENSITIVITY
h("10. Sensitivity and Injection-Dispersion Analysis", 1)
para("Launch never delivers the nominal orbit. Run injection-dispersion cases and re-run all "
     "figures of merit:")
bullet("Vary altitude +/-25 km, inclination +/-0.05-0.1 deg, RAAN by a few degrees (LTAN "
       "uncertainty), and epoch/phasing.")
bullet("Re-run coverage, revisit, contacts, beta/eclipse, and lifetime for each case.")
bullet("Report which requirements are sensitive vs. robust. Automate with STK Analyzer or a "
       "Python/STK-Engine parametric loop.")
para("This closes the Analysis (A) verification method and demonstrates the orbit meets "
     "requirements across the real injection envelope, not only at the nominal point.", italic=True)

# ============================================================ 11. AUTOMATION
h("11. Automation and Reproducibility", 1)
bullet("Drive STK from Python (agi.stk12 / STK Engine) or Connect/MATLAB so the study is fully "
       "scripted: build satellite -> configure HPOP -> run coverage -> export CSV/plots.")
bullet("Version-control scripts, input TLEs, and SpaceWeather files (this repository) for "
       "reproducible, auditable review-board results.")
bullet("Export a standard report pack: ground track, access tables, coverage FOM maps, "
       "beta/eclipse plots, lifetime curve, and dispersion summary.")

# ============================================================ 12. CHECKLIST
h("12. Deliverable Checklist", 1)
for item in [
    "Three orbit definitions (TLE / Keplerian / SSO-wizard) reconciled",
    "RAAN precession = +0.9856 deg/day confirmed (true SSO)",
    "HPOP force model documented (gravity, drag, SRP, third-body)",
    "Beta-angle and eclipse profile over a full year",
    "GS contact statistics vs. data-budget requirement",
    "Payload coverage: revisit time and % coverage for AOI",
    "Cluster/rideshare relative motion vs. CAS500-3",
    "Orbit lifetime and re-entry compliance",
    "Injection-dispersion sensitivity study",
    "Scripted, version-controlled, reproducible report pack",
]:
    doc.add_paragraph("[  ] " + item, style="List Bullet")

# ============================================================ PRIORITY
h("13. Highest-Priority Actions", 1)
numbered("Validate true sun-synchronicity (Section 5) before anything else; everything depends on it.")
numbered("Use HPOP with a realistic drag model (Section 4); at 550 km, SGP4/J2 alone will "
         "mislead coverage and lifetime.")
numbered("Bracket everything with min/max drag and solar-flux cases (Sections 4, 8, 10); a "
         "CubeSat without propulsion lives or dies by the atmosphere.")

doc.save(OUT)
print("Wrote", OUT)
