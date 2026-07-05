"""
Generates: KARIOS_SysML_System_Model.docx
Detailed description of the KARIOS SysML v2 system model with its five diagram
views (BDD, Requirements, Use Case, State Machine, Parametric) embedded.
Companion to KARIOS_system_model.sysml and KARIOS_sysml_diagrams.tex.

Run:  python build_sysml_doc.py
Requires: pip install python-docx  ;  the _sysml_pg-*.png views must exist.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "KARIOS_SysML_System_Model.docx"
PAGES = {
    1: ("Block Definition Diagram (bdd) — Structure",
        "The system decomposition. KARIOS_Mission composes the Space, Ground, and Launch "
        "segments; the KARIOS_Spacecraft composes the Bus and the two payloads (SIGINT primary, "
        "cybersecurity secondary); the Bus composes the six subsystems (OBC, EPS, ADCS, Comms "
        "with UHF + LoRaWAN, Structure, Thermal). Black diamonds denote composition."),
    2: ("Requirements Diagram (req) — traceability",
        "Mission requirements (MR) are refined into system requirements (SR-1..SR-5), each "
        "backed by a parametric constraint. Dashed open-arrow «satisfy» relationships link the "
        "requirements to the parts that satisfy them (EPS, Comms, the payloads), forming the "
        "basis of the Requirements Verification Matrix (RVM)."),
    3: ("Use Case Diagram (uc) — behaviour context",
        "The KARIOS system boundary with its four mission use cases (Collect SIGINT, Distribute "
        "Mission Data, Operate Spacecraft, Run Cyber Experiment) and the external actors "
        "(SIGINT Target, Ground Station, Mission Operator, Cyber Researcher)."),
    4: ("State Machine (stm) — operational states",
        "Four concurrent state regions that may overlap operationally: orbitState "
        "(Launched/Operational), solarState (Sunlight/Eclipse), payloadState (Idle/Collecting), "
        "and downloadState (Transmitting only within a ground-station access window)."),
    5: ("Parametric Diagram (par) — engineering budgets",
        "The mission budgets expressed as «constraint» blocks — power, data, link, and coverage "
        "— with their parameters bound to value properties. STK- and MATLAB-computed values "
        "(contact time, eclipse fraction, revisit time) are bound in through ModelCenter as an "
        "executable co-simulation."),
    6: ("Internal Block Diagram (ibd) — internal connections",
        "The internal wiring of the KARIOS_Spacecraft: EPS distributes power to the OBC, ADCS, "
        "Comms, and both payloads; the OBC data bus connects to Comms and receives payload data; "
        "the OBC issues commands to ADCS; and Comms drives the external UHF and LoRaWAN ports. "
        "Ports are shown as small squares and connectors as orthogonal item flows."),
    7: ("Activity Diagram (act) — Collect & Distribute Mission Data",
        "The top-level mission activity across four partitions (Payload, OBC, Comms, Ground "
        "Station): generate tasking, collect data, store on-board, and — when a ground station "
        "is in view — downlink the data frame for receipt and distribution. Object flows carry "
        "the Raw Data and Data Frame items; the decision node guards downlink on station access."),
}

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
INK2   = RGBColor(0x3F, 0x4B, 0x5E)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    return p


def make_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ---- title ----
tp = doc.add_heading("KARIOS CubeSat — SysML System Model", level=0)
for r in tp.runs:
    r.font.color.rgb = ACCENT
para("A SysML v2 model of the KAIST Intelligence & Reconnaissance Orbital Satellite (KARIOS), "
     "instantiated from the INCOSE SSWG CubeSat Reference Model. The authoritative model is the "
     "textual file KARIOS_system_model.sysml; this document describes it and renders its diagram "
     "views.", italic=True)

# ---- 1 organisation ----
h("1. Model Organisation", 1)
para("The model is authored in SysML v2 textual notation and organised into packages that map "
     "onto the four SysML pillars plus the mission context:")
make_table(
    ["Package", "Content", "SysML view"],
    [
        ["MissionContext", "External entities: space/Earth environment, launch provider, "
         "regulators, SIGINT target, ground stations, operator, cyber researcher.", "Domain / uc"],
        ["Interfaces", "Item, port and interface definitions (power, data, RF, command).", "ibd"],
        ["Structure", "Part definitions & compositions: mission, spacecraft, bus, subsystems, "
         "payloads, with internal connections.", "bdd / ibd"],
        ["Behavior", "Use cases, the Collect & Distribute Mission Data activity, and the "
         "operational state machine.", "uc / act / stm"],
        ["Analysis", "Orbit attributes and the power/data/link/coverage constraint blocks.", "par"],
        ["Requirements", "Mission & system requirements with satisfy/verify relationships.", "req"],
    ],
)
para("The textual model is the single source of truth: it is version-controllable, diff-able, "
     "and importable into SysML v2 tooling (e.g. the pilot implementation / SysIDE). The diagrams "
     "in Section 3 are views generated from — and kept consistent with — this model.")

# ---- 2 key elements ----
h("2. Key Model Elements", 1)
para("Structure — the KARIOS_Spacecraft is a 6U CubeSat with a Bus and two payloads. The Bus "
     "internal block diagram wires EPS power to OBC/ADCS/Comms and routes the OBC data bus to "
     "Comms and command lines to ADCS.")
para("Behaviour — the top-level activity CollectAndDistributeMissionData decomposes into "
     "generate-tasking, collect, store-on-board, and downlink actions; the state machine captures "
     "the concurrent orbit/solar/payload/download states used by the parametric simulation.")
para("Analysis — four constraint blocks encode the mission budgets. Their parameters are bound "
     "to structural value properties and to STK/MATLAB analysis results, making the model "
     "executable: requirements are verified by evaluating the constraints over a simulated orbit.")
para("Requirements — five system requirements each 'require' a constraint (power balance, data "
     "downlink, link margin, revisit) and are satisfied by specific parts, giving end-to-end "
     "traceability from mission need to design element to verification method.")

# ---- 3 diagrams ----
h("3. Diagram Views", 1)
missing = []
for n in sorted(PAGES):
    title, desc = PAGES[n]
    h(f"3.{n}  {title}", 2)
    para(desc)
    png = f"_sysml_pg-{n}.png"
    if os.path.exists(png):
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(png, width=Inches(6.6))
    else:
        missing.append(png)
        para(f"[view image {png} not found — rasterize KARIOS_sysml_diagrams.pdf]", italic=True)

# ---- 4 references ----
h("4. References", 1)
for ref in [
    "KARIOS_system_model.sysml — the authoritative SysML v2 model (this document describes it).",
    "KARIOS_sysml_diagrams.tex / .pdf — rendered diagram views.",
    "INCOSE SSWG CubeSat Reference Model (CRM); OMG SysML v2 specification.",
    "D. Kaslow et al., MBSE applied to the RAX CubeSat mission (IEEE Aerospace).",
    "KARIOS mission draft (KAIST Intelligence & Reconnaissance Orbital Satellite).",
]:
    doc.add_paragraph(ref, style="List Bullet")

doc.save(OUT)
print("Wrote", OUT, "| missing views:", missing if missing else "none")
