"""
Generates: KARIOS_Verification_Process.docx
KARIOS CubeSat Mission Verification Process Proposal, including a schematic
block-diagram rendered as a vertical flow of shaded stage boxes, and the
compiled LaTeX/TikZ diagram embedded as a figure if the PNG is present.

Run:  python build_verification_doc.py
Requires: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "KARIOS_Verification_Process.docx"
DIAGRAM_PNG = "_diagram_preview-1.png"   # rendered from the LaTeX diagram

ACCENT   = RGBColor(0x1F, 0x4E, 0x79)
INK2     = RGBColor(0x3F, 0x4B, 0x5E)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
# method colors
C_TEST   = RGBColor(0xC9, 0x76, 0x1F)
C_ANAL   = RGBColor(0x2F, 0x74, 0xC0)
C_REV    = RGBColor(0x7C, 0x53, 0xA5)
C_INSP   = RGBColor(0x1F, 0x8F, 0x86)

SHADE_LIGHT = "F1F4F9"
SHADE_KEY   = "E4EEF8"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def set_cell_shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def stage_box(title, desc, key=False, chips=None):
    """One block of the schematic flow: a shaded, bordered 1x1 table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_shade(cell, SHADE_KEY if key else SHADE_LIGHT)
    cell.width = Inches(6.2)
    # title
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(11)
    # description
    if desc:
        pd = cell.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rd = pd.add_run(desc)
        rd.font.size = Pt(9)
        rd.font.color.rgb = INK2
    # method/model chips line
    if chips:
        pc = cell.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, (label, col) in enumerate(chips):
            if i:
                pc.add_run("    ").font.size = Pt(9)
            rr = pc.add_run(label)
            rr.bold = True
            rr.font.size = Pt(9)
            if col:
                rr.font.color.rgb = col
    return t


def arrow():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("▼")   # down triangle
    r.font.color.rgb = ACCENT
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ============================================================ TITLE
tp = doc.add_heading("KARIOS CubeSat — Mission Verification Process Proposal", level=0)
for r in tp.runs:
    r.font.color.rgb = ACCENT
para("A structured, model-based verification programme demonstrating that the KARIOS CubeSat "
     "meets every mission requirement — from conceptual design through in-orbit commissioning "
     "(ECSS-E-ST-10, MBSE / SysML).", italic=True)

para("Model strategy: EM → EQM → PFM      |      Review gates: full ECSS set "
     "(MDR, PRR, PDR, CDR, QR, AR, FRR, CRR, ELR)", bold=True)

# ============================================================ 1
h("1. Purpose and Scope", 1)
para("This proposal defines the verification process for the KARIOS CubeSat mission, spanning "
     "initial conceptual design to in-orbit commissioning. Its purpose is to demonstrate that the "
     "delivered KARIOS spacecraft achieves all mission objectives — validating the design and "
     "verifying that the integrated system satisfies its requirements before launch.")
para("The process is built on the ECSS space-engineering life-cycle (ECSS-E-ST-10), environmental "
     "testing standards, and CubeSat design specifications for structural and dimensional "
     "compliance. It is tailored for a small-satellite programme with constrained budget and "
     "schedule, and is executed within a Model-Based Systems Engineering (MBSE) framework using "
     "SysML, per the INCOSE Space Systems Working Group approach.")

# ============================================================ 2
h("2. Verification Methods", 1)
para("Every requirement is closed by one or more of four verification methods. Each method is "
     "allocated in the Verification Control Document (VCD) and traced back to a mission or system "
     "requirement.")
make_table(
    ["Method", "Description", "Primary use in KARIOS"],
    [
        ["Test (T)", "Physical measurement on representative or flight hardware.",
         "Structure, thermal, EMC/EMI, functional performance, deployment."],
        ["Analysis (A)", "Analytical / computational prediction where test is impractical.",
         "Orbital mechanics & coverage (STK), structural stress, thermal model, link & power budgets."],
        ["Review of Design (RoD)", "Inspection of design documentation and heritage evidence.",
         "Interface control, heritage / COTS component reuse justification."],
        ["Inspection (I)", "Direct visual and dimensional examination.",
         "Dimensional & mass-envelope compliance to the CubeSat Design Specification."],
    ],
)
para("Orbital dynamics and coverage are closed primarily by Analysis (A): the pre-launch STK "
     "simulation of KARIOS orbital dynamics, sun-synchronicity, lighting, ground-station contacts, "
     "coverage, and lifetime forms the analytical evidence for this method.", italic=True)

# ============================================================ 3
h("3. Model Philosophy", 1)
para("KARIOS follows a three-model strategy adapted from the ECSS Proto-Flight Model (PFM) "
     "approach — chosen because it delivers full qualification confidence within the cost and "
     "schedule limits of a CubeSat programme by qualifying and flying representative hardware.")
make_table(
    ["Model", "Role", "Test level"],
    [
        ["Engineering Model (EM)",
         "Electrically & mechanically representative unit for early functional and interface "
         "verification. Flight-representative PCBs/panels; components need not be flight-grade.",
         "Functional / development"],
        ["Engineering Qualification Model (EQM)",
         "Combined functional and qualification-level environmental testing (vibration, shock, "
         "thermal-vacuum at qualification margins) validating the design envelope.",
         "Qualification"],
        ["Proto-Flight Model (PFM)",
         "The flight article, acceptance-tested at proto-flight levels (qualification amplitude, "
         "acceptance duration). Delivered for integration, launch campaign, and flight.",
         "Proto-flight / acceptance"],
    ],
)

# ============================================================ 4  LIFECYCLE
h("4. Life-Cycle Phases and Review Gates", 1)
para("The verification programme is executed against the ECSS-E-ST-10 life-cycle. Each phase "
     "closes at a formal review gate that dispositions the accumulated verification evidence.")
make_table(
    ["Phase", "Definition", "Review gate(s)"],
    [
        ["0 / A", "Mission analysis, needs identification & feasibility", "MDR, PRR"],
        ["B", "Preliminary definition — verification planning baselined", "PDR"],
        ["C", "Detailed definition — methods allocated & executed", "CDR"],
        ["D", "Qualification & production — model + environmental campaign", "QR, AR"],
        ["E", "Operations — launch, LEOP, in-orbit demonstration & commissioning", "FRR, CRR"],
        ["F", "Disposal — end-of-life & re-entry", "ELR"],
    ],
)

# ============================================================ 5  DIAGRAM (flow)
h("5. Verification Process — Schematic Block Diagram", 1)
para("The flow reads top-to-bottom against the ECSS phases and their gate reviews. Requirements "
     "drive verification planning; planning allocates the four methods; methods are executed "
     "across the model campaign and environmental test campaign; results are dispositioned in the "
     "VCD and closed at each review gate. Non-conformances (NCR / RID) feed back to design for "
     "re-verification.")

stage_box("Mission & System Requirements  [Phase 0 / A / B]",
          "Captured in the SysML MBSE model  →  Requirements Verification Matrix (RVM).", key=True)
arrow()
stage_box("Verification Planning  [Phase B → PDR]",
          "Verification Plan authored; each requirement allocated a method, level, and model. "
          "Baseline of the Verification Control Document (VCD).")
arrow()
stage_box("Verification Methods — allocation & execution  [Phase C / D]",
          "Four methods allocated per requirement and executed:", key=True,
          chips=[("Test (T)", C_TEST), ("Analysis (A)", C_ANAL),
                 ("Review (RoD)", C_REV), ("Inspection (I)", C_INSP)])
arrow()
stage_box("Model Campaign  [Phase C / D]",
          "Verification distributed across the three-model strategy:",
          chips=[("EM", ACCENT), ("→  EQM", ACCENT), ("→  PFM", ACCENT)])
arrow()
stage_box("Environmental Test Campaign  [Phase D → QR / AR]",
          "Functional/DITL · Sine & random vibration · Mechanical shock · Thermal-vacuum & "
          "balance · EMC/EMI · Mass properties · Deployment test · Bake-out/outgassing.", key=True)
arrow()
stage_box("Verification Close-out  [QR · AR]",
          "Results dispositioned in the VCD; Verification Control Board closes each requirement "
          "with objective evidence. Every RVM item marked verified.")
arrow()
stage_box("Flight Readiness → Launch → LEOP  [Phase D → E · FRR]",
          "Pre-ship review, launch campaign, separation, and Launch & Early-Orbit Phase acquisition.",
          key=True)
arrow()
stage_box("In-Orbit Demonstration & Commissioning  [Phase E · CRR]",
          "On-orbit performance confirmed against requirements (IOD); mission validation and "
          "hand-over to routine operations.", key=True)

fb = doc.add_paragraph()
fb.alignment = WD_ALIGN_PARAGRAPH.CENTER
rfb = fb.add_run("↺  Feedback: any non-conformance (NCR / RID) at test or review routes back "
                 "to Design & the MBSE model for corrective action and re-verification.")
rfb.italic = True
rfb.font.size = Pt(9)
rfb.font.color.rgb = C_TEST

# embed the rendered LaTeX diagram, if available
if os.path.exists(DIAGRAM_PNG):
    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Figure 1 — Schematic block diagram of the KARIOS verification process "
                     "(rendered from KARIOS_verification_block_diagram.tex).")
    rc.italic = True
    rc.font.size = Pt(9)
    rc.font.color.rgb = INK2
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(DIAGRAM_PNG, width=Inches(6.3))

# ============================================================ 6  ENV DETAIL
h("6. Environmental Test Campaign", 1)
para("The environmental campaign qualifies the design (EQM) and acceptance-tests the flight "
     "article (PFM) to launch-provider and CubeSat-deployer requirements:")
for item in [
    "Functional / Day-in-the-Life (DITL) — end-to-end system performance and mode transitions.",
    "Sine & random vibration — to qualification levels (EQM) and proto-flight/acceptance (PFM).",
    "Mechanical shock — separation and deployment shock environments.",
    "Thermal-vacuum & thermal balance — hot/cold cycling and thermal-model correlation.",
    "EMC / EMI — self-compatibility and emissions.",
    "Mass properties — mass, CoG, and moments of inertia versus allocation.",
    "Deployment test — antennas, solar panels, and separation mechanisms.",
    "Bake-out / outgassing — contamination control before delivery.",
]:
    bullet(item)

# ============================================================ 7  STANDARDS
h("7. Governing Standards and Close-out", 1)
for item in [
    "ECSS-E-ST-10 — system engineering life-cycle and verification framework (phases 0–F, gate reviews).",
    "ECSS-E-ST-10-02 / -03 — verification process and testing requirements.",
    "CubeSat Design Specification (CDS) — dimensional, mass, and mechanical-envelope compliance (Inspection).",
    "Launch-provider & deployer ICD — vibration/shock qualification levels and interface constraints.",
    "ISO 19683 — small-spacecraft design-qualification and acceptance test methods.",
]:
    bullet(item)
para("The programme is complete when every requirement in the RVM is closed with objective "
     "evidence in the VCD, all non-conformances are dispositioned, and in-orbit commissioning "
     "confirms mission performance — formally validating the KARIOS design and verifying the "
     "system against its requirements.")

doc.save(OUT)
print("Wrote", OUT, "(diagram embedded)" if os.path.exists(DIAGRAM_PNG) else "(no diagram PNG found)")
